import streamlit as st
import os
import tempfile
import json
import time
from moviepy import VideoFileClip
import google.generativeai as genai
from groq import Groq
from pydub import AudioSegment
import math
import nltk
from nltk.tokenize import sent_tokenize, word_tokenize
from nltk.corpus import stopwords
from nltk.probability import FreqDist
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import io
import shutil
import sys
from datetime import datetime
# Importações para resolver o problema de ScriptRunContext
from streamlit.runtime.scriptrunner import add_script_run_ctx, get_script_run_ctx
import logging

# Configurar o logging para suprimir avisos específicos
logger = logging.getLogger('streamlit')
logger.setLevel(logging.ERROR)

# Pegar o contexto atual do Streamlit
ctx = get_script_run_ctx()

# Verificar se FFmpeg está instalado
def check_ffmpeg():
    try:
        # Verificar se ffmpeg está no PATH
        ffmpeg_path = shutil.which('ffmpeg')
        if ffmpeg_path is None:
            # Verificar em caminhos específicos do Streamlit Cloud
            additional_paths = [
                "/usr/bin",
                "/usr/local/bin",
                "/app/.apt/usr/bin",  # Caminho específico do Streamlit Cloud
                "/home/appuser/.apt/usr/bin"  # Outro caminho possível no Streamlit Cloud
            ]
            
            # Verificar cada caminho manualmente
            for path in additional_paths:
                potential_path = os.path.join(path, 'ffmpeg')
                if os.path.exists(potential_path) and os.access(potential_path, os.X_OK):
                    # Adicionar ao PATH para futuras verificações
                    os.environ["PATH"] = path + os.pathsep + os.environ["PATH"]
                    return True
            return False
        return True
    except Exception as e:
        st.error(f"Erro ao verificar FFmpeg: {str(e)}")
        return False

# Configuração inicial da página
st.set_page_config(
    page_title="Processador de Áudio e Transcrição",
    page_icon="🎙️",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Estilo CSS personalizado
st.markdown("""
<style>
    /* Cores e tema */
    :root {
        --primary: #6C63FF;
        --primary-light: #8F88FF;
        --primary-dark: #5046E5;
        --secondary: #FF6584;
        --accent: #43BCCD;
        --background: #F8F9FD;
        --card: #FFFFFF;
        --text: #333333;
        --text-light: #666666;
        --success: #4CAF50;
        --warning: #FF9800;
        --error: #F44336;
        --border-radius: 12px;
        --shadow: 0 4px 12px rgba(0, 0, 0, 0.08);
    }

    /* Estilos gerais */
    .stApp {
        background-color: var(--background);
    }
    
    h1, h2, h3, h4, h5, h6 {
        font-family: 'Poppins', sans-serif;
        font-weight: 600;
    }
    
    p, div {
        font-family: 'Inter', sans-serif;
    }
    
    /* Cabeçalho principal */
    .main-header {
        display: flex;
        align-items: center;
        margin-bottom: 2rem;
        padding: 1.5rem;
        background: linear-gradient(135deg, var(--primary), var(--primary-dark));
        border-radius: var(--border-radius);
        color: white;
        box-shadow: var(--shadow);
    }
    
    .main-header-icon {
        font-size: 2.5rem;
        margin-right: 1rem;
    }
    
    .main-header-text h1 {
        font-size: 2rem;
        margin: 0;
        padding: 0;
    }
    
    .main-header-text p {
        margin: 0.5rem 0 0 0;
        opacity: 0.9;
    }
    
    /* Cartões de seção */
    .section-card {
        background-color: var(--card);
        border-radius: var(--border-radius);
        padding: 1.5rem;
        margin-bottom: 1.5rem;
        box-shadow: var(--shadow);
        border-top: 5px solid var(--primary);
    }
    
    .section-header {
        display: flex;
        align-items: center;
        margin-bottom: 1.5rem;
    }
    
    .section-number {
        display: flex;
        align-items: center;
        justify-content: center;
        width: 36px;
        height: 36px;
        background-color: var(--primary);
        color: white;
        border-radius: 50%;
        font-weight: bold;
        margin-right: 1rem;
    }
    
    .section-title {
        font-size: 1.5rem;
        color: var(--text);
        margin: 0;
    }
    
    /* Mensagens e alertas */
    .info-message, .success-message, .warning-message, .error-message {
        padding: 1rem;
        border-radius: var(--border-radius);
        margin: 1rem 0;
        display: flex;
        align-items: center;
    }
    
    .info-message {
        background-color: rgba(67, 188, 205, 0.1);
        border-left: 4px solid var(--accent);
    }
    
    .success-message {
        background-color: rgba(76, 175, 80, 0.1);
        border-left: 4px solid var(--success);
    }
    
    .warning-message {
        background-color: rgba(255, 152, 0, 0.1);
        border-left: 4px solid var(--warning);
    }
    
    .error-message {
        background-color: rgba(244, 67, 54, 0.1);
        border-left: 4px solid var(--error);
    }
    
    .message-icon {
        font-size: 1.5rem;
        margin-right: 1rem;
    }
    
    /* Botões personalizados */
    .custom-button {
        display: inline-flex;
        align-items: center;
        justify-content: center;
        padding: 0.5rem 1.5rem;
        background-color: var(--primary);
        color: white;
        border: none;
        border-radius: 50px;
        font-weight: 600;
        cursor: pointer;
        transition: all 0.3s ease;
        text-decoration: none;
        margin: 0.5rem 0;
    }
    
    .custom-button:hover {
        background-color: var(--primary-dark);
        box-shadow: 0 4px 8px rgba(0, 0, 0, 0.1);
        transform: translateY(-2px);
    }
    
    .custom-button-secondary {
        background-color: white;
        color: var(--primary);
        border: 2px solid var(--primary);
    }
    
    .custom-button-secondary:hover {
        background-color: var(--primary-light);
        color: white;
    }
    
    .custom-button-icon {
        margin-right: 0.5rem;
        font-size: 1.2rem;
    }
    
    /* Upload de arquivo personalizado */
    .file-uploader {
        border: 2px dashed #ccc;
        border-radius: var(--border-radius);
        padding: 2rem;
        text-align: center;
        transition: all 0.3s ease;
        cursor: pointer;
    }
    
    .file-uploader:hover {
        border-color: var(--primary);
        background-color: rgba(108, 99, 255, 0.05);
    }
    
    .file-uploader-icon {
        font-size: 3rem;
        color: var(--primary);
        margin-bottom: 1rem;
    }
    
    /* Barra de progresso personalizada */
    .progress-container {
        margin: 1.5rem 0;
    }
    
    .progress-label {
        display: flex;
        justify-content: space-between;
        margin-bottom: 0.5rem;
    }
    
    .progress-bar {
        height: 8px;
        background-color: #e0e0e0;
        border-radius: 4px;
        overflow: hidden;
    }
    
    .progress-bar-fill {
        height: 100%;
        background: linear-gradient(90deg, var(--primary), var(--primary-light));
        border-radius: 4px;
        transition: width 0.3s ease;
    }
    
    /* Animações */
    @keyframes pulse {
        0% { transform: scale(1); }
        50% { transform: scale(1.05); }
        100% { transform: scale(1); }
    }
    
    .pulse {
        animation: pulse 2s infinite;
    }
    
    /* Sidebar personalizada */
    .css-1d391kg, .css-1lcbmhc {
        background-color: var(--card);
    }
    
    .sidebar-header {
        padding: 1rem;
        background: linear-gradient(135deg, var(--primary-dark), var(--primary));
        color: white;
        border-radius: var(--border-radius);
        margin-bottom: 1rem;
        text-align: center;
    }
    
    /* Responsividade */
    @media (max-width: 768px) {
        .main-header {
            flex-direction: column;
            text-align: center;
        }
        
        .main-header-icon {
            margin-right: 0;
            margin-bottom: 1rem;
        }
    }
    
    /* Estilizando elementos nativos do Streamlit */
    .stTextInput > div > div > input {
        border-radius: var(--border-radius);
    }
    
    .stButton > button {
        border-radius: 50px;
        padding: 0.5rem 2rem;
        font-weight: 600;
        background-color: var(--primary);
        color: white;
        border: none;
        transition: all 0.3s ease;
    }
    
    .stButton > button:hover {
        background-color: var(--primary-dark);
        box-shadow: 0 4px 8px rgba(0, 0, 0, 0.1);
        transform: translateY(-2px);
    }
    
    .stProgress > div > div > div {
        background: linear-gradient(90deg, var(--primary), var(--primary-light));
    }
    
    /* Estilizando o player de áudio/vídeo */
    .stAudio > div, .stVideo > div {
        border-radius: var(--border-radius);
        overflow: hidden;
        box-shadow: var(--shadow);
    }
    
    /* Estilizando áreas de texto */
    .stTextArea > div > div > textarea {
        border-radius: var(--border-radius);
        border: 1px solid #e0e0e0;
    }
    
    /* Estilizando o download button */
    .stDownloadButton > button {
        border-radius: 50px;
        padding: 0.5rem 2rem;
        font-weight: 600;
        background-color: var(--success);
        color: white;
        border: none;
        transition: all 0.3s ease;
        display: inline-flex;
        align-items: center;
    }
    
    .stDownloadButton > button:hover {
        background-color: #3d8b40;
        box-shadow: 0 4px 8px rgba(0, 0, 0, 0.1);
        transform: translateY(-2px);
    }
    
    /* Estilizando tooltips */
    .tooltip {
        position: relative;
        display: inline-block;
        cursor: help;
    }
    
    .tooltip .tooltiptext {
        visibility: hidden;
        width: 200px;
        background-color: #333;
        color: #fff;
        text-align: center;
        border-radius: 6px;
        padding: 5px;
        position: absolute;
        z-index: 1;
        bottom: 125%;
        left: 50%;
        margin-left: -100px;
        opacity: 0;
        transition: opacity 0.3s;
    }
    
    .tooltip:hover .tooltiptext {
        visibility: visible;
        opacity: 1;
    }
    
    /* Estilizando o spinner */
    .loading-spinner {
        display: flex;
        justify-content: center;
        align-items: center;
        margin: 2rem 0;
    }
    
    .spinner {
        border: 4px solid rgba(0, 0, 0, 0.1);
        width: 36px;
        height: 36px;
        border-radius: 50%;
        border-left-color: var(--primary);
        animation: spin 1s linear infinite;
    }
    
    @keyframes spin {
        0% { transform: rotate(0deg); }
        100% { transform: rotate(360deg); }
    }
    
    /* Estilizando o status de cada etapa */
    .step-status {
        display: flex;
        align-items: center;
        margin-bottom: 0.5rem;
    }
    
    .step-status-icon {
        width: 24px;
        height: 24px;
        border-radius: 50%;
        display: flex;
        align-items: center;
        justify-content: center;
        margin-right: 0.5rem;
        font-weight: bold;
        font-size: 0.8rem;
    }
    
    .step-status-pending {
        background-color: #e0e0e0;
        color: #666;
    }
    
    .step-status-active {
        background-color: var(--primary);
        color: white;
    }
    
    .step-status-completed {
        background-color: var(--success);
        color: white;
    }
    
    .step-status-text {
        font-size: 0.9rem;
    }
</style>
""", unsafe_allow_html=True)

# Criar diretório temporário para arquivos
temp_dir = os.path.join(tempfile.gettempdir(), 'streamlit_audio_processor')
os.makedirs(temp_dir, exist_ok=True)

# Configuração para suprimir avisos de ScriptRunContext
def configure_script_run_context():
    """
    Configura o contexto de execução do Streamlit para evitar avisos de ScriptRunContext
    """
    try:
        # Suprime avisos específicos de logging
        logging.getLogger('streamlit.runtime.scriptrunner.script_run_context').setLevel(logging.ERROR)
        
        # Se houver threads em execução, adiciona o contexto a elas
        import threading
        for thread in threading.enumerate():
            if thread != threading.current_thread():
                try:
                    add_script_run_ctx(thread, ctx)
                except:
                    pass
    except Exception as e:
        # Não exibe erro para o usuário, apenas ignora silenciosamente
        pass

# Chamar a função para configurar o contexto
configure_script_run_context()

# Inicialização de variáveis de sessão
if 'audio_path' not in st.session_state:
    st.session_state.audio_path = None
if 'transcription_path' not in st.session_state:
    st.session_state.transcription_path = None
if 'transcription_text' not in st.session_state:
    st.session_state.transcription_text = None
if 'docx_path' not in st.session_state:
    st.session_state.docx_path = None
if 'docx_bytes' not in st.session_state:
    st.session_state.docx_bytes = None
if 'config_saved' not in st.session_state:
    st.session_state.config_saved = False
if 'current_step' not in st.session_state:
    st.session_state.current_step = 1
if 'steps_completed' not in st.session_state:
    st.session_state.steps_completed = []
if 'processing' not in st.session_state:
    st.session_state.processing = False

# Função para verificar se as configurações estão salvas
def check_config():
    config_path = os.path.join(temp_dir, 'config.json')
    try:
        if os.path.exists(config_path):
            with open(config_path, 'r') as f:
                config = json.load(f)
            return True
        return False
    except:
        return False

# Função para carregar as configurações
def load_config():
    config_path = os.path.join(temp_dir, 'config.json')
    with open(config_path, 'r') as f:
        return json.load(f)

# Classe para extração de áudio
class AudioExtractor:
    @staticmethod
    def extract_audio(video_path, output_path):
        try:
            video = VideoFileClip(video_path)
            audio = video.audio
            audio.write_audiofile(output_path)
            return output_path
        except Exception as e:
            st.error(f"Erro ao extrair áudio: {str(e)}")
            if "ffmpeg" in str(e).lower():
                st.error("Certifique-se de que o FFmpeg está instalado corretamente.")
            return None

# Classe para transcrição de áudio
class GroqTranscriber:
    def __init__(self, api_key):
        self.client = Groq(api_key=api_key)
        self.MAX_FILE_SIZE = 25 * 1024 * 1024  # 25 MB em bytes
        self.MODEL_ID = "whisper-large-v3-turbo"
        
    def get_file_size(self, file_path):
        """Retorna o tamanho do arquivo em bytes"""
        return os.path.getsize(file_path)
    
    def split_audio(self, audio_path, chunk_duration=300000):  # 300000 ms = 5 minutos
        """
        Divide o áudio em partes menores
        Retorna: Lista com os caminhos dos arquivos temporários
        """
        try:
            audio = AudioSegment.from_file(audio_path)
            duration = len(audio)
            chunks = []
            
            # Calcula o número de chunks necessários baseado no tamanho do arquivo
            file_size = self.get_file_size(audio_path)
            if file_size > self.MAX_FILE_SIZE:
                num_chunks = math.ceil(file_size / (self.MAX_FILE_SIZE * 0.95))  # 95% do limite para margem de segurança
                chunk_duration = duration / num_chunks
            
            # Duração mínima de um chunk em ms (0.01 segundos = 10ms)
            MIN_CHUNK_DURATION = 100  # 100ms para ter uma margem de segurança
            
            # Divide o áudio
            for i in range(0, duration, int(chunk_duration)):
                end_point = min(i + chunk_duration, duration)
                
                # Verifica se o chunk tem a duração mínima
                if (end_point - i) < MIN_CHUNK_DURATION:
                    # Se o chunk for muito curto, pule-o
                    continue
                    
                chunk = audio[i:end_point]
                chunk_path = os.path.join(temp_dir, f"temp_chunk_{i}.mp3")
                chunk.export(chunk_path, format="mp3")
                chunks.append(chunk_path)
                
            return chunks
        except Exception as e:
            st.error(f"Erro ao dividir áudio: {str(e)}")
            return []
    
    def transcribe_chunk(self, chunk_path):
        """Transcreve um único chunk de áudio usando a API do Groq"""
        try:
            # Verificar se o arquivo existe e tem tamanho mínimo
            if not os.path.exists(chunk_path) or os.path.getsize(chunk_path) < 100:  # 100 bytes mínimos
                st.warning(f"Arquivo de chunk muito pequeno ou inexistente: {chunk_path}")
                return ""
                
            with open(chunk_path, "rb") as audio_file:
                # Usa a API do Groq para transcrição
                response = self.client.audio.transcriptions.create(
                    model=self.MODEL_ID,
                    file=audio_file,
                    language="pt"
                )
                return response.text
        except Exception as e:
            error_msg = str(e)
            if "Audio file is too short" in error_msg:
                st.warning(f"Chunk {chunk_path} ignorado por ser muito curto")
            else:
                st.error(f"Erro ao transcrever chunk {chunk_path}: {error_msg}")
            return ""
    
    def transcribe(self, audio_path, output_path=None, progress_callback=None):
        """
        Transcreve um arquivo de áudio, dividindo-o se necessário
        """
        try:
            if output_path is None:
                output_path = os.path.join(temp_dir, "transcricao.txt")
                
            # Verifica se o arquivo precisa ser dividido
            if self.get_file_size(audio_path) <= self.MAX_FILE_SIZE:
                # Se não precisar dividir, transcreve diretamente
                if progress_callback:
                    progress_callback(0.5)
                transcription = self.transcribe_chunk(audio_path)
                with open(output_path, "w", encoding="utf-8") as f:
                    f.write(transcription)
                if progress_callback:
                    progress_callback(1.0)
                return transcription, output_path
            
            # Se precisar dividir, processa em chunks
            chunks = self.split_audio(audio_path)
            if not chunks:
                raise Exception("Falha ao dividir o áudio em partes")
                
            full_transcription = ""
            
            # Transcreve cada chunk - aqui pode estar usando threads internamente
            # Usar uma abordagem sequencial para evitar problemas com ScriptRunContext
            for i, chunk_path in enumerate(chunks, 1):
                if progress_callback:
                    # Garantir que as chamadas de callback tenham o contexto do Streamlit
                    if ctx:
                        # Segurar o contexto atual do Streamlit
                        add_script_run_ctx(progress_callback, ctx)
                    progress_callback(i / len(chunks))
                transcription = self.transcribe_chunk(chunk_path)
                full_transcription += transcription + "\n"
                
                # Remove o arquivo temporário
                try:
                    os.remove(chunk_path)
                except:
                    pass
            
            # Salva a transcrição completa
            with open(output_path, "w", encoding="utf-8") as f:
                f.write(full_transcription)
            
            return full_transcription, output_path
            
        except Exception as e:
            st.error(f"Erro durante a transcrição: {str(e)}")
            return None, None

# Classe para resumo e geração de documento
class SummaryGenerator:
    def __init__(self, api_key):
        genai.configure(api_key=api_key)
        self.model = genai.GenerativeModel('gemini-1.5-flash')
        
        # Baixar recursos necessários do NLTK
        try:
            nltk.data.find('tokenizers/punkt')
        except LookupError:
            nltk.download('punkt')
        
        try:
            nltk.data.find('corpora/stopwords')
        except LookupError:
            nltk.download('stopwords')
            
        # Tentar baixar o recurso punkt_tab (pode não estar disponível)
        try:
            nltk.data.find('tokenizers/punkt_tab')
        except LookupError:
            try:
                nltk.download('punkt_tab')
            except:
                # Se não conseguir baixar o punkt_tab, usaremos o método padrão
                st.warning("Não foi possível baixar o recurso NLTK 'punkt_tab'. Usando método alternativo para tokenização.")
        
        # Configurar o tokenizador - com fallback seguro
        try:
            # Tentar carregar o tokenizador para português
            self.tokenizer = nltk.data.load('tokenizers/punkt/portuguese.pickle')
        except:
            # Se falhar, criar um tokenizador simples
            class SimpleTokenizer:
                def tokenize(self, text):
                    # Um tokenizador de sentenças simples baseado em pontuação
                    import re
                    return re.split(r'(?<=[.!?])\s+', text)
            self.tokenizer = SimpleTokenizer()
    
    def preprocessar_texto(self, texto):
        # Tokenização de sentenças e palavras
        # Usando o tokenizador personalizado em vez de sent_tokenize
        sentencas = self.tokenizer.tokenize(texto)
        palavras = word_tokenize(texto.lower())
        
        # Remover stopwords
        stop_words = set(stopwords.words('portuguese'))
        palavras_filtradas = [palavra for palavra in palavras if palavra.isalnum() and palavra not in stop_words]
        
        # Análise de frequência de palavras
        freq_dist = FreqDist(palavras_filtradas)
        palavras_mais_comuns = freq_dist.most_common(10)
        
        return sentencas, palavras_filtradas, palavras_mais_comuns

    def gerar_resumo_profissional(self, conteudo, progress_callback=None):
        if progress_callback:
            # Adicionando contexto do Streamlit ao callback
            if ctx:
                add_script_run_ctx(progress_callback, ctx)
            progress_callback(0.3)
            
        sentencas, palavras_filtradas, palavras_mais_comuns = self.preprocessar_texto(conteudo)

        prompt = f"""
        Como professor de português, analise o conteúdo do arquivo em anexo e faça um texto destacando os principais pontos, com um tom profissional.

        Conteúdo original:
        {conteudo}

        Análise prévia:
        - Número de sentenças: {len(sentencas)}
        - Palavras-chave mais frequentes: {', '.join([palavra for palavra, _ in palavras_mais_comuns])}

        O resumo deve:
        1. Ter um tom profissional e formal
        2. Destacar os pontos-chave do texto original
        3. Ser estruturado em tópicos claros
        4. Usar a seguinte formatação:
           - Tópicos principais sem marcadores
           - Subtópicos com • (ex: • Subtópico)
           - Sub-subtópicos com dois espaços e • (ex:   • Sub-subtópico)
        5. Manter uma estrutura hierárquica clara
        6. NÃO usar asteriscos (**) ou hashtags (###) na formatação

        Por favor, formate o resumo de maneira clara e legível, seguindo estritamente as regras de formatação acima.
        """

        if progress_callback:
            progress_callback(0.5)
            
        try:
            resposta = self.model.generate_content(prompt)
            
            if progress_callback:
                progress_callback(1.0)
                
            return resposta.text
        except Exception as e:
            st.error(f"Erro ao gerar resumo: {str(e)}")
            return None

    def salvar_como_doc(self, conteudo, caminho_arquivo=None):
        if caminho_arquivo is None:
            caminho_arquivo = os.path.join(temp_dir, "resumo.docx")
            
        documento = Document()

        # Definir estilos
        estilos = {
            'Normal': {
                'nome': 'Calibri',
                'tamanho': 11,
                'alinhamento': WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            },
            'Titulo': {
                'nome': 'Calibri',
                'tamanho': 12,
                'negrito': True,
                'alinhamento': WD_PARAGRAPH_ALIGNMENT.LEFT
            },
            'Marcador1': {
                'nome': 'Calibri',
                'tamanho': 11,
                'alinhamento': WD_PARAGRAPH_ALIGNMENT.LEFT,
                'estilo': 'List Bullet'
            },
            'Marcador2': {
                'nome': 'Calibri',
                'tamanho': 11,
                'alinhamento': WD_PARAGRAPH_ALIGNMENT.LEFT,
                'estilo': 'List Bullet 2'
            }
        }

        # Criar ou atualizar estilos
        for nome, config in estilos.items():
            if nome in documento.styles:
                estilo = documento.styles[nome]
            else:
                estilo = documento.styles.add_style(nome, WD_STYLE_TYPE.PARAGRAPH)
            
            estilo.font.name = config['nome']
            estilo.font.size = Pt(config['tamanho'])
            if 'negrito' in config:
                estilo.font.bold = config['negrito']
            if 'estilo' in config:
                estilo.base_style = documento.styles[config['estilo']]

        # Processar o conteúdo
        linhas = conteudo.split('\n')
        nivel_atual = 0

        for linha in linhas:
            linha = linha.strip()
            if not linha:
                continue

            # Remover asteriscos e hashtags
            linha = linha.replace('*', '').replace('#', '')

            if not linha.startswith('•'):
                # Tópico principal
                p = documento.add_paragraph(linha, style='Titulo')
                nivel_atual = 0
            elif linha.startswith('•'):
                # Marcador de primeiro nível
                p = documento.add_paragraph(linha.lstrip('• '), style='Marcador1')
                nivel_atual = 1
            elif linha.startswith('  •'):
                # Marcador de segundo nível
                p = documento.add_paragraph(linha.lstrip(' •'), style='Marcador2')
                nivel_atual = 2
            else:
                # Texto normal
                p = documento.add_paragraph(linha, style='Normal')

            p.alignment = estilos[p.style.name]['alinhamento']

        try:
            # Salvar o documento
            documento.save(caminho_arquivo)
            
            # Retornar também os bytes para download
            file_stream = io.BytesIO()
            documento.save(file_stream)
            file_stream.seek(0)
            
            return file_stream.getvalue(), caminho_arquivo
        except Exception as e:
            st.error(f"Erro ao salvar documento: {str(e)}")
            return None, None

# Função para exibir o status de cada etapa
def display_step_status(step_number, step_name):
    status_class = "step-status-pending"
    status_text = "Pendente"
    
    if step_number == st.session_state.current_step:
        status_class = "step-status-active"
        status_text = "Em andamento"
    elif step_number in st.session_state.steps_completed:
        status_class = "step-status-completed"
        status_text = "Concluído"
    
    html = f"""
    <div class="step-status">
        <div class="step-status-icon {status_class}">{step_number}</div>
        <div class="step-status-text">{step_name}: <strong>{status_text}</strong></div>
    </div>
    """
    st.markdown(html, unsafe_allow_html=True)

# Função para exibir mensagem de erro personalizada
def display_error(message):
    st.markdown(f"""
    <div class="error-message">
        <div class="message-icon">❌</div>
        <div>{message}</div>
    </div>
    """, unsafe_allow_html=True)

# Função para exibir mensagem de sucesso personalizada
def display_success(message):
    st.markdown(f"""
    <div class="success-message">
        <div class="message-icon">✅</div>
        <div>{message}</div>
    </div>
    """, unsafe_allow_html=True)

# Função para exibir mensagem de informação personalizada
def display_info(message):
    st.markdown(f"""
    <div class="info-message">
        <div class="message-icon">ℹ️</div>
        <div>{message}</div>
    </div>
    """, unsafe_allow_html=True)

# Função para exibir mensagem de aviso personalizada
def display_warning(message):
    st.markdown(f"""
    <div class="warning-message">
        <div class="message-icon">⚠️</div>
        <div>{message}</div>
    </div>
    """, unsafe_allow_html=True)

# Função para exibir barra de progresso personalizada
def display_progress(progress, label=""):
    st.markdown(f"""
    <div class="progress-container">
        <div class="progress-label">
            <span>{label}</span>
            <span>{int(progress * 100)}%</span>
        </div>
        <div class="progress-bar">
            <div class="progress-bar-fill" style="width: {int(progress * 100)}%;"></div>
        </div>
    </div>
    """, unsafe_allow_html=True)

# Função para exibir spinner personalizado
def display_spinner():
    st.markdown("""
    <div class="loading-spinner">
        <div class="spinner"></div>
    </div>
    """, unsafe_allow_html=True)

# Verificar se FFmpeg está instalado
ffmpeg_available = check_ffmpeg()

# Sidebar para configurações
with st.sidebar:
    st.markdown('<div class="sidebar-header"><h3>⚙️ Configurações</h3></div>', unsafe_allow_html=True)
    
    # Exibir status das etapas
    st.markdown("### Status do Processo")
    display_step_status(1, "Upload e Extração")
    display_step_status(2, "Transcrição")
    display_step_status(3, "Resumo e Documento")
    
    st.markdown("---")
    
    st.markdown("### Configurações de API")
    
    google_api_key = st.text_input("Google API Key (Gemini)", type="password")
    groq_api_key = st.text_input("Groq API Key", type="password", value="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")
    
    if st.button("Salvar Configurações"):
        # Criar arquivo de configuração
        config = {
            "GOOGLE_API_KEY": google_api_key,
            "GROQ_API_KEY": groq_api_key
        }
        
        config_path = os.path.join(temp_dir, 'config.json')
        with open(config_path, 'w', encoding='utf-8') as f:
            json.dump(config, f)
        
        st.session_state.config_saved = True
        display_success("Configurações salvas com sucesso!")
    
    st.markdown("---")
    
    # Informações do aplicativo
    st.markdown("### Sobre o Aplicativo")
    st.markdown("""
    Este aplicativo permite processar vídeos para:
    - Extrair áudio
    - Transcrever conteúdo
    - Gerar resumos profissionais
    - Criar documentos formatados
    """)
    
    # Data e hora atual
    now = datetime.now()
    st.markdown(f"**Data atual:** {now.strftime('%d/%m/%Y %H:%M')}")

# Cabeçalho principal
st.markdown("""
<div class="main-header">
    <div class="main-header-icon">🎙️</div>
    <div class="main-header-text">
        <h1>Processador de Áudio e Transcrição</h1>
        <p>Extraia áudio de vídeos, transcreva e gere resumos profissionais em formato DOCX.</p>
    </div>
</div>
""", unsafe_allow_html=True)

# Alerta de FFmpeg
if not ffmpeg_available:
    display_warning("""
    FFmpeg não encontrado! O FFmpeg é necessário para processamento de áudio e vídeo.
    
    **Instruções de instalação:**
    
    - **Windows**: Baixe do [site oficial](https://ffmpeg.org/download.html) ou use Chocolatey: `choco install ffmpeg`
    - **Mac**: Use Homebrew: `brew install ffmpeg`
    - **Linux**: Use apt: `sudo apt install ffmpeg`
    
    Após instalar, reinicie o aplicativo.
    """)

# Seção 1: Upload de vídeo e extração de áudio
st.markdown("""
<div class="section-card">
    <div class="section-header">
        <div class="section-number">1</div>
        <div class="section-title">Upload de Vídeo e Extração de Áudio</div>
    </div>
""", unsafe_allow_html=True)

# Área de upload de arquivo
st.markdown("""
<div class="file-uploader">
    <div class="file-uploader-icon">📁</div>
    <h3>Arraste e solte seu arquivo de vídeo aqui</h3>
    <p>Ou clique para selecionar um arquivo</p>
</div>
""", unsafe_allow_html=True)

video_file = st.file_uploader("Selecione um arquivo de vídeo", type=['mp4', 'avi', 'mov', 'mkv'])

if video_file is not None:
    # Salvar o arquivo de vídeo temporariamente
    video_path = os.path.join(temp_dir, video_file.name)
    with open(video_path, 'wb') as f:
        f.write(video_file.getbuffer())
    
    # Exibir o vídeo
    st.video(video_path)
    
    col1, col2 = st.columns([1, 1])
    
    with col1:
        st.markdown(f"""
        <div class="info-message">
            <div class="message-icon">📋</div>
            <div>
                <strong>Detalhes do arquivo:</strong><br>
                Nome: {video_file.name}<br>
                Tamanho: {video_file.size / (1024*1024):.2f} MB
            </div>
        </div>
        """, unsafe_allow_html=True)
    
    extract_button = st.button("🔊 Extrair Áudio", key="extract_audio")
    
    if extract_button:
        if not ffmpeg_available:
            display_error("FFmpeg não está instalado. Por favor, instale o FFmpeg para continuar.")
        else:
            st.session_state.processing = True
            
            with st.spinner("Extraindo áudio do vídeo..."):
                # Criar nome de arquivo baseado no nome original
                file_name = os.path.splitext(video_file.name)[0]
                audio_path = os.path.join(temp_dir, f"{file_name}.mp3")
                
                # Exibir progresso
                progress_placeholder = st.empty()
                progress_placeholder.markdown('<div class="progress-container"><div class="progress-label"><span>Extraindo áudio...</span><span>0%</span></div><div class="progress-bar"><div class="progress-bar-fill" style="width: 0%;"></div></div></div>', unsafe_allow_html=True)
                
                # Extrair áudio
                extractor = AudioExtractor()
                audio_path = extractor.extract_audio(video_path, audio_path)
                
                # Atualizar progresso
                for i in range(10):
                    progress_placeholder.markdown(f'<div class="progress-container"><div class="progress-label"><span>Extraindo áudio...</span><span>{(i+1)*10}%</span></div><div class="progress-bar"><div class="progress-bar-fill" style="width: {(i+1)*10}%;"></div></div></div>', unsafe_allow_html=True)
                    time.sleep(0.1)
                
                if audio_path:
                    # Salvar o caminho do áudio na sessão
                    st.session_state.audio_path = audio_path
                    st.session_state.current_step = 2
                    st.session_state.steps_completed.append(1)
                    
                    display_success("Áudio extraído com sucesso!")
                    
                    # Exibir player de áudio
                    try:
                        audio_file = open(audio_path, 'rb')
                        audio_bytes = audio_file.read()
                        st.audio(audio_bytes, format='audio/mp3')
                        audio_file.close()
                    except Exception as e:
                        display_error(f"Erro ao exibir áudio: {str(e)}")
                else:
                    display_error("Falha ao extrair áudio. Verifique se o FFmpeg está instalado corretamente.")
            
            st.session_state.processing = False

st.markdown("</div>", unsafe_allow_html=True)

# Seção 2: Transcrição de áudio
st.markdown("""
<div class="section-card">
    <div class="section-header">
        <div class="section-number">2</div>
        <div class="section-title">Transcrição de Áudio</div>
    </div>
""", unsafe_allow_html=True)

if st.session_state.audio_path and os.path.exists(st.session_state.audio_path):
    display_info(f"Áudio pronto para transcrição: {os.path.basename(st.session_state.audio_path)}")
    
    transcribe_button = st.button("🔤 Transcrever Áudio", key="transcribe_audio")
    
    if transcribe_button:
        # Verificar se as configurações estão salvas
        if not check_config():
            display_warning("Por favor, configure as chaves de API na barra lateral antes de continuar.")
        else:
            st.session_state.processing = True
            config = load_config()
            
            # Criar barra de progresso
            progress_placeholder = st.empty()
            status_text = st.empty()
            
            # Configurar contexto do Streamlit antes de iniciar a transcrição
            configure_script_run_context()
            
            def update_progress(progress):
                progress_placeholder.markdown(f'<div class="progress-container"><div class="progress-label"><span>Transcrevendo áudio...</span><span>{int(progress * 100)}%</span></div><div class="progress-bar"><div class="progress-bar-fill" style="width: {int(progress * 100)}%;"></div></div></div>', unsafe_allow_html=True)
                status_text.markdown(f"<p>Processando... Por favor, aguarde.</p>", unsafe_allow_html=True)
            
            with st.spinner("Transcrevendo áudio..."):
                # Iniciar transcrição
                transcriber = GroqTranscriber(api_key=config["GROQ_API_KEY"])
                transcription, output_path = transcriber.transcribe(
                    st.session_state.audio_path,
                    progress_callback=update_progress
                )
                
                if transcription and output_path:
                    # Salvar o caminho e o texto da transcrição na sessão
                    st.session_state.transcription_path = output_path
                    st.session_state.transcription_text = transcription
                    st.session_state.current_step = 3
                    st.session_state.steps_completed.append(2)
                    
                    display_success("Transcrição concluída com sucesso!")
                    
                    # Exibir prévia da transcrição
                    st.markdown("### Prévia da Transcrição")
                    st.text_area("Transcrição", transcription, height=200)
                    
                    # Botão para baixar a transcrição
                    st.download_button(
                        label="📝 Baixar Transcrição TXT",
                        data=transcription,
                        file_name=f"{os.path.splitext(os.path.basename(st.session_state.audio_path))[0]}_transcricao.txt",
                        mime="text/plain"
                    )
                else:
                    display_error("Falha na transcrição. Verifique os logs para mais detalhes.")
            
            st.session_state.processing = False
elif st.session_state.audio_path:
    display_error(f"O arquivo de áudio não foi encontrado: {st.session_state.audio_path}")
    st.session_state.audio_path = None
else:
    display_info("Extraia o áudio de um vídeo primeiro.")

st.markdown("</div>", unsafe_allow_html=True)

# Seção 3: Resumo e Geração de Documento
st.markdown("""
<div class="section-card">
    <div class="section-header">
        <div class="section-number">3</div>
        <div class="section-title">Resumo e Geração de Documento</div>
    </div>
""", unsafe_allow_html=True)

if st.session_state.transcription_text:
    display_info("Transcrição pronta para resumo")
    
    summarize_button = st.button("📄 Gerar Resumo e Documento", key="generate_summary")
    
    if summarize_button:
        # Verificar se as configurações estão salvas
        if not check_config():
            display_warning("Por favor, configure as chaves de API na barra lateral antes de continuar.")
        else:
            st.session_state.processing = True
            config = load_config()
            
            # Configurar contexto do Streamlit antes de iniciar o processo
            configure_script_run_context()
            
            # Criar barra de progresso
            progress_placeholder = st.empty()
            status_text = st.empty()
            
            def update_progress(progress):
                progress_placeholder.markdown(f'<div class="progress-container"><div class="progress-label"><span>Gerando resumo...</span><span>{int(progress * 100)}%</span></div><div class="progress-bar"><div class="progress-bar-fill" style="width: {int(progress * 100)}%;"></div></div></div>', unsafe_allow_html=True)
                status_text.markdown(f"<p>Processando... Por favor, aguarde.</p>", unsafe_allow_html=True)
            
            with st.spinner("Gerando resumo e documento..."):
                # Criar nome de arquivo baseado no nome original
                file_name = os.path.splitext(os.path.basename(st.session_state.transcription_path)) if st.session_state.transcription_path else ("resumo",)
                file_name = file_name[0]
                docx_path = os.path.join(temp_dir, f"{file_name}.docx")
                
                # Gerar resumo e documento
                generator = SummaryGenerator(api_key=config["GOOGLE_API_KEY"])
                
                # Gerar resumo
                resumo_profissional = generator.gerar_resumo_profissional(
                    st.session_state.transcription_text,
                    progress_callback=update_progress
                )
                
                if resumo_profissional:
                    # Salvar como documento
                    docx_bytes, docx_path = generator.salvar_como_doc(resumo_profissional, docx_path)
                    
                    if docx_bytes and docx_path:
                        # Salvar o caminho do documento na sessão
                        st.session_state.docx_path = docx_path
                        st.session_state.docx_bytes = docx_bytes
                        st.session_state.steps_completed.append(3)
                        
                        display_success("Documento gerado com sucesso!")
                        
                        # Exibir prévia do resumo
                        st.markdown("### Prévia do Resumo")
                        st.text_area("Resumo", resumo_profissional, height=200)
                        
                        # Botão para download
                        col1, col2 = st.columns([1, 1])
                        
                        with col1:
                            st.download_button(
                                label="📥 Baixar Documento DOCX",
                                data=docx_bytes,
                                file_name=os.path.basename(docx_path),
                                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                            )
                        
                        with col2:
                            st.markdown("""
                            <div style="background-color: #E8F5E9; padding: 1rem; border-radius: 0.5rem; text-align: center;">
                                <h3 style="margin-top: 0;">✅ Processo Concluído!</h3>
                                <p>Todos os passos foram concluídos com sucesso.</p>
                            </div>
                            """, unsafe_allow_html=True)
                    else:
                        display_error("Falha ao gerar o documento DOCX.")
                else:
                    display_error("Falha ao gerar o resumo.")
            
            st.session_state.processing = False
elif st.session_state.transcription_path and os.path.exists(st.session_state.transcription_path):
    # Tentar carregar o texto da transcrição do arquivo
    try:
        with open(st.session_state.transcription_path, 'r', encoding='utf-8') as f:
            st.session_state.transcription_text = f.read()
        display_info("Transcrição carregada do arquivo")
    except Exception as e:
        display_error(f"Erro ao carregar transcrição do arquivo: {str(e)}")
        st.session_state.transcription_path = None
else:
    display_info("Transcreva um áudio primeiro.")

st.markdown("</div>", unsafe_allow_html=True)

# Rodapé
st.markdown("""
<div style="margin-top: 3rem; padding: 1rem; background-color: #f5f5f5; border-radius: 0.5rem; text-align: center;">
    <p>Desenvolvido com ❤️ usando Streamlit, Groq e Google Gemini</p>
    <p style="font-size: 0.8rem; color: #666;">© 2025 - Todos os direitos reservados</p>
</div>
""", unsafe_allow_html=True)

# Limpar arquivos temporários antigos (mais de 1 hora)
def clean_temp_files():
    try:
        current_time = time.time()
        for filename in os.listdir(temp_dir):
            file_path = os.path.join(temp_dir, filename)
            if os.path.isfile(file_path):
                # Se o arquivo tem mais de 1 hora
                if current_time - os.path.getmtime(file_path) > 3600:
                    try:
                        os.remove(file_path)
                    except:
                        pass
    except:
        pass

# Executar limpeza de arquivos temporários
clean_temp_files()

